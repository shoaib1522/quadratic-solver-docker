from docx import Document
from docx.shared import Inches

def create_report():
    doc = Document()
    
    # Title
    doc.add_heading('Assignment 02: Hosting and Load Balancing a Web Application Using Docker Containers', 0)
    doc.add_paragraph('Course: DS 584 - Cloud Computing, FCIT, University of the Punjab')
    doc.add_paragraph('Instructor: Junaid Abdullah')
    doc.add_paragraph('Semester: Spring 2026')
    doc.add_page_break()
    
    # Milestone 1
    doc.add_heading('Milestone 01: Develop a Sample Web Application', level=1)
    doc.add_paragraph('Objective: Develop a quadratic equation solver web application using HTML, JQuery, Bootstrap, and Flask (server-side). Use AJAX to submit inputs and display solution without postback.')
    doc.add_paragraph('[SCREENSHOT 1: Web application showing form and solution for a=1, b=-5, c=6 with server IP]')
    doc.add_paragraph('Source Code:')
    doc.add_paragraph('app.py:')
    with open('app.py', 'r') as f:
        doc.add_paragraph(f.read())
    doc.add_paragraph('templates/index.html:')
    with open('templates/index.html', 'r') as f:
        doc.add_paragraph(f.read())
    doc.add_page_break()
    
    # Milestone 2
    doc.add_heading('Milestone 02: Host your Web Application in a Docker Container', level=1)
    doc.add_paragraph('Objective: Prepare a Docker image of the web application and run two instances. Verify accessibility from host OS browser.')
    doc.add_paragraph('[SCREENSHOT 2: Browser accessing http://localhost showing web app]')
    doc.add_paragraph('[SCREENSHOT 3: Output of docker ps command showing 3 containers]')
    doc.add_paragraph('Source Code:')
    doc.add_paragraph('Dockerfile:')
    with open('Dockerfile', 'r') as f:
        doc.add_paragraph(f.read())
    doc.add_paragraph('docker-compose.yml:')
    with open('docker-compose.yml', 'r') as f:
        doc.add_paragraph(f.read())
    doc.add_page_break()
    
    # Milestone 3
    doc.add_heading('Milestone 03: Deploy Nginx as a Proxy Server in a Docker Container', level=1)
    doc.add_paragraph('Objective: Use Nginx as a reverse proxy and load balancer to distribute traffic between the two application containers.')
    doc.add_paragraph('[SCREENSHOT 4: Browser via nginx showing load balancing (changing server IPs)]')
    doc.add_paragraph('[SCREENSHOT 5: docker ps output with nginx, app1, app2 running]')
    doc.add_paragraph('Nginx Configuration:')
    doc.add_paragraph('nginx/nginx.conf:')
    with open('nginx/nginx.conf', 'r') as f:
        doc.add_paragraph(f.read())
    doc.add_page_break()
    
    # Milestone 4
    doc.add_heading('Milestone 04: Synthetic Workload Generation', level=1)
    doc.add_paragraph('Objective: Use httperf to generate synthetic workload to the Nginx container and profile CPU usage. Generate two graphs: average response time and CPU utilization.')
    doc.add_paragraph('[SCREENSHOT 6: Graph 1 - Average Response Time: graphs/response_time_graph.png]')
    doc.add_paragraph('[SCREENSHOT 7: Graph 2 - CPU Utilization: graphs/cpu_utilization_graph.png]')
    doc.add_paragraph('Source Code:')
    doc.add_paragraph('cpu_profiler.py:')
    with open('cpu_profiler.py', 'r') as f:
        doc.add_paragraph(f.read())
    doc.add_paragraph('generate_graphs.py:')
    with open('generate_graphs.py', 'r') as f:
        doc.add_paragraph(f.read())
    
    # Save
    doc.save('Assignment_Report.docx')
    print('Report saved as Assignment_Report.docx')

if __name__ == '__main__':
    create_report()