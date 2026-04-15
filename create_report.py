import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level=0):
    doc.add_heading(text, level=level)

def add_paragraph(doc, text, style=None):
    p = doc.add_paragraph(text)
    if style:
        p.style = style
    return p

def add_code_block(doc, code, filename):
    p = doc.add_paragraph()
    run = p.add_run(f'Code: {filename}')
    run.bold = True
    doc.add_paragraph(code, style='Intense Quote')

def add_screenshot_placeholder(doc, description):
    p = doc.add_paragraph()
    run = p.add_run(f'[SCREENSHOT PLACEHOLDER: {description}]')
    run.italic = True
    run.font.color.rgb = (0, 100, 0)  # dark green
    doc.add_paragraph()  # empty line

def main():
    doc = Document()
    # Set font for whole document
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    add_heading(doc, 'Assignment 02: Hosting and Load Balancing a Web Application Using Docker Containers', level=0)
    add_paragraph(doc, 'Course: DS 584 - Cloud Computing, FCIT, University of the Punjab')
    add_paragraph(doc, 'Instructor: Junaid Abdullah')
    add_paragraph(doc, 'Semester: Spring 2026')
    doc.add_page_break()
    
    # Milestone 1
    add_heading(doc, 'Milestone 01: Develop a Sample Web Application', level=1)
    add_paragraph(doc, 'Objective: Develop a quadratic equation solver web application using HTML, JQuery, Bootstrap, and Flask (server-side). Use AJAX to submit inputs and display solution without postback.')
    add_screenshot_placeholder(doc, 'Screenshot of the web application showing the form and an example solution (e.g., a=1, b=-5, c=6) with server IP displayed.')
    add_paragraph(doc, 'Source Code:')
    # app.py
    with open('app.py', 'r', encoding='utf-8') as f:
        app_code = f.read()
    add_code_block(doc, app_code, 'app.py')
    # templates/index.html
    with open(os.path.join('templates', 'index.html'), 'r', encoding='utf-8') as f:
        index_code = f.read()
    add_code_block(doc, index_code, 'templates/index.html')
    doc.add_page_break()
    
    # Milestone 2
    add_heading(doc, 'Milestone 02: Host your Web Application in a Docker Container', level=1)
    add_paragraph(doc, 'Objective: Prepare a Docker image of the web application and run two instances. Verify accessibility from host OS browser.')
    add_screenshot_placeholder(doc, 'Screenshot of the browser accessing the web application via http://localhost (showing the form and a solution).')
    add_paragraph(doc, 'Command output:')
    add_paragraph(doc, '$ docker ps', style='Intense Quote')
    # We'll add placeholder for docker ps output; could fetch actual but we'll leave placeholder.
    add_paragraph(doc, '[Docker PS Output Placeholder - show three containers: nginx, app1, app2]')
    add_paragraph(doc, 'Source Code:')
    with open('Dockerfile', 'r', encoding='utf-8') as f:
        dockerfile_code = f.read()
    add_code_block(doc, dockerfile_code, 'Dockerfile')
    with open('docker-compose.yml', 'r', encoding='utf-8') as f:
        compose_code = f.read()
    add_code_block(doc, compose_code, 'docker-compose.yml')
    doc.add_page_break()
    
    # Milestone 3
    add_heading(doc, 'Milestone 03: Deploy Nginx as a Proxy Server in a Docker Container', level=1)
    add_paragraph(doc, 'Objective: Use Nginx as a reverse proxy and load balancer to distribute traffic between the two application containers.')
    add_screenshot_placeholder(doc, 'Screenshot of the browser accessing the application via the Nginx proxy (http://localhost) showing load balancing (note changing server IPs/hostnames on refresh).')
    add_paragraph(doc, 'Nginx Configuration:')
    with open(os.path.join('nginx', 'nginx.conf'), 'r', encoding='utf-8') as f:
        nginx_code = f.read()
    add_code_block(doc, nginx_code, 'nginx/nginx.conf')
    add_paragraph(doc, 'Source Code (same as Milestone 2):')
    add_paragraph(doc, 'Dockerfile and docker-compose.yml as above.')
    add_screenshot_placeholder(doc, 'Screenshot of docker ps showing three containers running (nginx, app1, app2).')
    doc.add_page_break()
    
    # Milestone 4
    add_heading(doc, 'Milestone 04: Synthetic Workload Generation', level=1)
    add_paragraph(doc, 'Objective: Use httperf to generate synthetic workload to the Nginx container and profile CPU usage. Generate two graphs: average response time and CPU utilization.')
    add_screenshot_placeholder(doc, 'Graph 1: Average Response Time (from httperf output).')
    add_paragraph(doc, 'Graph 2: CPU Utilization Over Time (from CPU profiler).')
    # Indicate where to place images
    add_paragraph(doc, '[Insert Graph 1: response_time_graph.png here]', style='Intense Quote')
    add_paragraph(doc, '[Insert Graph 2: cpu_utilization_graph.png here]', style='Intense Quote')
    add_paragraph(doc, 'Source Code:')
    with open('cpu_profiler.py', 'r', encoding='utf-8') as f:
        cpu_code = f.read()
    add_code_block(doc, cpu_code, 'cpu_profiler.py')
    with open('generate_graphs.py', 'r', encoding='utf-8') as f:
        graph_code = f.read()
    add_code_block(doc, graph_code, 'generate_graphs.py')
    doc.add_page_break()
    
    # Conclusion
    add_heading(doc, 'Conclusion', level=1)
    add_paragraph(doc, 'All milestones have been successfully completed. The web application is containerized, load balanced via Nginx, and performance evaluated with synthetic workload.')
    
    # Save document
    output_path = 'Assignment_Report.docx'
    doc.save(output_path)
    print(f'Document saved to {output_path}')

if __name__ == '__main__':
    main()